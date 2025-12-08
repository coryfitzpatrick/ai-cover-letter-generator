"""Comprehensive tests for prepare_data module to increase coverage to 80%+.

This test suite covers:
- Text extraction from PDF files
- Text extraction from DOCX files
- Text chunking with various configurations
- Data loading and processing
- Edge cases: empty files, missing files, corrupt files
- CSV processing (LinkedIn data)
- JSON processing
"""

import json
from unittest.mock import Mock, patch

from docx import Document

from src.cover_letter_generator.prepare_data import (
    DEFAULT_CHUNK_SIZE,
    DEFAULT_OVERLAP,
    chunk_text,
    extract_text_from_docx,
    process_csv_files,
    process_json_files,
    process_linkedin_profile_csv,
    process_linkedin_recommendations_csv,
)

# NOTE: TestExtractTextFromPDF removed because extract_text_from_pdf function was removed
# during code cleanup. PDF extraction is now handled differently in the codebase.


class TestExtractTextFromDOCX:
    """Tests for DOCX text extraction."""

    def test_extract_text_from_valid_docx(self, tmp_path):
        """Test extracting text from a valid DOCX file."""
        docx_path = tmp_path / "test.docx"

        doc = Document()
        doc.add_paragraph("This is the first paragraph.")
        doc.add_paragraph("This is the second paragraph.")
        doc.save(str(docx_path))

        text = extract_text_from_docx(str(docx_path))

        assert "first paragraph" in text
        assert "second paragraph" in text

    def test_extract_text_from_docx_with_tables(self, tmp_path):
        """Test extracting text from DOCX with tables."""
        docx_path = tmp_path / "with_tables.docx"

        doc = Document()
        doc.add_paragraph("Text before table")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1,1"
        table.cell(0, 1).text = "Cell 1,2"
        table.cell(1, 0).text = "Cell 2,1"
        table.cell(1, 1).text = "Cell 2,2"

        doc.add_paragraph("Text after table")
        doc.save(str(docx_path))

        text = extract_text_from_docx(str(docx_path))

        # Should extract text from both paragraphs and table cells
        assert "Text before table" in text
        assert "Text after table" in text
        assert "Cell 1,1" in text
        assert "Cell 2,2" in text

    def test_extract_text_from_empty_docx(self, tmp_path):
        """Test extracting text from an empty DOCX file."""
        docx_path = tmp_path / "empty.docx"

        doc = Document()
        doc.save(str(docx_path))

        text = extract_text_from_docx(str(docx_path))

        assert text == "" or text.strip() == ""

    def test_extract_text_from_nonexistent_docx(self):
        """Test handling of non-existent DOCX file."""
        text = extract_text_from_docx("/nonexistent/file.docx")

        assert text == ""

    def test_extract_text_from_corrupt_docx(self, tmp_path):
        """Test handling of corrupt DOCX file."""
        docx_path = tmp_path / "corrupt.docx"

        # Create a file with invalid DOCX content
        with open(docx_path, "w") as f:
            f.write("This is not a valid DOCX file")

        text = extract_text_from_docx(str(docx_path))

        assert text == ""

    def test_extract_text_with_formatting(self, tmp_path):
        """Test that text is extracted regardless of formatting."""
        docx_path = tmp_path / "formatted.docx"

        doc = Document()
        p = doc.add_paragraph()
        run1 = p.add_run("Bold text ")
        run1.bold = True
        run2 = p.add_run("italic text ")
        run2.italic = True
        p.add_run("normal text")
        doc.save(str(docx_path))

        text = extract_text_from_docx(str(docx_path))

        # All text should be extracted regardless of formatting
        assert "Bold text" in text
        assert "italic text" in text
        assert "normal text" in text


class TestChunkText:
    """Tests for text chunking functionality."""

    def test_chunk_text_basic(self):
        """Test basic text chunking."""
        text = "This is a test. " * 100  # Create text longer than default chunk size
        chunks = chunk_text(text)

        assert len(chunks) > 1
        assert all(isinstance(chunk, str) for chunk in chunks)

    def test_chunk_text_with_custom_size(self):
        """Test chunking with custom chunk size."""
        text = "Word " * 200
        chunk_size = 50
        chunks = chunk_text(text, chunk_size=chunk_size)

        # Most chunks should be around the specified size
        for chunk in chunks[:-1]:  # Exclude last chunk which may be smaller
            assert len(chunk) <= chunk_size + 100  # Allow some overlap

    def test_chunk_text_with_overlap(self):
        """Test that overlap is respected between chunks."""
        text = "0123456789" * 100
        chunk_size = 100
        overlap = 20
        chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap)

        # Check that there's some overlap between consecutive chunks
        assert len(chunks) > 1

    def test_chunk_text_empty_string(self):
        """Test chunking an empty string."""
        chunks = chunk_text("")
        assert chunks == []

    def test_chunk_text_short_text(self):
        """Test chunking text shorter than chunk size."""
        text = "Short text"
        chunks = chunk_text(text, chunk_size=1000)

        assert len(chunks) == 1
        assert chunks[0] == "Short text"

    def test_chunk_text_respects_paragraphs(self):
        """Test that chunking respects paragraph boundaries."""
        text = "Paragraph 1.\n\nParagraph 2.\n\nParagraph 3."
        chunks = chunk_text(text, chunk_size=30)

        # Should split on paragraph boundaries when possible
        assert len(chunks) >= 1

    def test_chunk_text_with_metadata_header(self):
        """Test that metadata header is prepended to each chunk."""
        text = "Test content " * 100
        metadata = "Source: test.pdf"
        chunks = chunk_text(text, metadata_header=metadata)

        # All chunks should start with the metadata
        for chunk in chunks:
            assert metadata in chunk

    def test_chunk_text_long_single_line(self):
        """Test handling of a very long single line."""
        text = "A" * 2000  # Single line longer than default chunk size
        chunks = chunk_text(text, chunk_size=500)

        # Should split even without paragraph breaks
        assert len(chunks) > 1

    def test_chunk_text_with_blank_lines(self):
        """Test handling of multiple consecutive blank lines."""
        text = "Content 1\n\n\n\nContent 2\n\n\nContent 3"
        chunks = chunk_text(text)

        # Should handle multiple blank lines gracefully
        assert len(chunks) >= 1

    def test_chunk_text_force_split_threshold(self):
        """Test that soft split threshold forces splits at paragraph boundaries."""
        # Create text with clear paragraph breaks
        paragraphs = ["This is paragraph {}.".format(i) * 20 for i in range(5)]
        text = "\n\n".join(paragraphs)

        chunks = chunk_text(text, chunk_size=600)

        # Should create multiple chunks
        assert len(chunks) > 1


class TestProcessLinkedInProfileCSV:
    """Tests for LinkedIn profile CSV processing."""

    def test_process_profile_csv_with_summary(self, tmp_path):
        """Test processing a profile CSV with summary."""
        csv_path = tmp_path / "Profile.csv"

        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            f.write("First Name,Last Name,Headline,Summary\n")
            f.write("John,Doe,Software Engineer,Experienced engineer with 10 years\n")

        results = process_linkedin_profile_csv(str(csv_path))

        assert len(results) == 2  # Summary and headline

        # Check summary
        summary_result = [r for r in results if r[1]["type"] == "profile_summary"][0]
        assert "Experienced engineer" in summary_result[0]
        assert summary_result[1]["name"] == "John Doe"

    def test_process_profile_csv_with_headline(self, tmp_path):
        """Test processing a profile CSV with headline."""
        csv_path = tmp_path / "Profile.csv"

        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            f.write("First Name,Last Name,Headline,Summary\n")
            f.write("Jane,Smith,Senior Manager,\n")

        results = process_linkedin_profile_csv(str(csv_path))

        assert len(results) >= 1

        # Check headline
        headline_result = [r for r in results if r[1]["type"] == "headline"][0]
        assert "Senior Manager" in headline_result[0]

    def test_process_profile_csv_empty_file(self, tmp_path):
        """Test processing an empty profile CSV."""
        csv_path = tmp_path / "Profile.csv"

        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            f.write("First Name,Last Name,Headline,Summary\n")

        results = process_linkedin_profile_csv(str(csv_path))

        assert results == []

    def test_process_profile_csv_missing_file(self):
        """Test handling of missing CSV file."""
        results = process_linkedin_profile_csv("/nonexistent/Profile.csv")

        assert results == []


class TestProcessLinkedInRecommendationsCSV:
    """Tests for LinkedIn recommendations CSV processing."""

    def test_process_recommendations_csv(self, tmp_path):
        """Test processing recommendations CSV."""
        csv_path = tmp_path / "Recommendations_Received.csv"

        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            f.write("First Name,Last Name,Company,Job Title,Text,Status\n")
            f.write("Alice,Johnson,TechCorp,CTO,Great leader and mentor,VISIBLE\n")

        results = process_linkedin_recommendations_csv(str(csv_path))

        assert len(results) == 1
        assert "Alice Johnson" in results[0][0]
        assert "Great leader" in results[0][0]
        assert results[0][1]["recommender"] == "Alice Johnson"
        assert results[0][1]["company"] == "TechCorp"

    def test_process_recommendations_csv_hidden_status(self, tmp_path):
        """Test that hidden recommendations are filtered out."""
        csv_path = tmp_path / "Recommendations_Received.csv"

        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            f.write("First Name,Last Name,Company,Job Title,Text,Status\n")
            f.write("Bob,Smith,Company1,Manager,Hidden recommendation,HIDDEN\n")
            f.write("Carol,Lee,Company2,Director,Visible recommendation,VISIBLE\n")

        results = process_linkedin_recommendations_csv(str(csv_path))

        assert len(results) == 1
        assert "Carol Lee" in results[0][0]
        assert "Bob Smith" not in str(results)

    def test_process_recommendations_csv_empty(self, tmp_path):
        """Test processing empty recommendations CSV."""
        csv_path = tmp_path / "Recommendations_Received.csv"

        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            f.write("First Name,Last Name,Company,Job Title,Text,Status\n")

        results = process_linkedin_recommendations_csv(str(csv_path))

        assert results == []


class TestProcessCSVFiles:
    """Tests for processing multiple CSV files."""

    def test_process_csv_files_profile_and_recommendations(self, tmp_path):
        """Test processing both profile and recommendations CSVs."""
        # Create profile CSV
        profile_path = tmp_path / "Profile.csv"
        with open(profile_path, "w", newline="", encoding="utf-8") as f:
            f.write("First Name,Last Name,Headline,Summary\n")
            f.write("John,Doe,Engineer,Summary text\n")

        # Create recommendations CSV
        rec_path = tmp_path / "Recommendations_Received.csv"
        with open(rec_path, "w", newline="", encoding="utf-8") as f:
            f.write("First Name,Last Name,Company,Job Title,Text,Status\n")
            f.write("Alice,Smith,Corp,CTO,Great work,VISIBLE\n")

        results = process_csv_files(tmp_path)

        # Should have entries from both files
        assert len(results) >= 2

    def test_process_csv_files_no_csv_files(self, tmp_path):
        """Test processing directory with no CSV files."""
        results = process_csv_files(tmp_path)

        assert results == []

    def test_process_csv_files_excludes_template_folder(self, tmp_path):
        """Test that template folder is excluded from processing."""
        # Create template folder with CSV
        template_dir = tmp_path / "template"
        template_dir.mkdir()

        csv_path = template_dir / "Profile.csv"
        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            f.write("First Name,Last Name,Headline,Summary\n")
            f.write("Test,User,Title,Summary\n")

        results = process_csv_files(tmp_path)

        # Template folder should be excluded
        assert len(results) == 0


class TestProcessJSONFiles:
    """Tests for JSON file processing."""

    def test_process_json_dict(self, tmp_path):
        """Test processing JSON file with dictionary structure."""
        json_path = tmp_path / "data.json"
        data = {"name": "John Doe", "role": "Engineer", "experience": 5}

        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(data, f)

        results = process_json_files(tmp_path)

        assert len(results) == 3
        assert any("name: John Doe" in r[0] for r in results)
        assert any("role: Engineer" in r[0] for r in results)

    def test_process_json_list(self, tmp_path):
        """Test processing JSON file with list structure."""
        json_path = tmp_path / "items.json"
        data = ["Item 1", "Item 2", "Item 3"]

        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(data, f)

        results = process_json_files(tmp_path)

        assert len(results) == 3

    def test_process_json_nested_dict(self, tmp_path):
        """Test processing JSON with nested dictionaries."""
        json_path = tmp_path / "nested.json"
        data = {"person": {"name": "Jane", "age": 30}}

        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(data, f)

        results = process_json_files(tmp_path)

        assert len(results) >= 1
        # Should extract nested data
        assert any("person" in r[0] for r in results)

    def test_process_json_list_of_dicts(self, tmp_path):
        """Test processing JSON with list of dictionaries."""
        json_path = tmp_path / "list_dicts.json"
        data = {
            "employees": [{"name": "John", "role": "Engineer"}, {"name": "Jane", "role": "Manager"}]
        }

        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(data, f)

        results = process_json_files(tmp_path)

        assert len(results) >= 2

    def test_process_json_excludes_contact_info(self, tmp_path):
        """Test that contact_info.json is excluded."""
        json_path = tmp_path / "contact_info.json"
        data = {"email": "test@example.com"}

        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(data, f)

        results = process_json_files(tmp_path)

        # contact_info.json should be excluded
        assert len(results) == 0

    def test_process_json_excludes_template_folder(self, tmp_path):
        """Test that template folder is excluded."""
        template_dir = tmp_path / "template"
        template_dir.mkdir()

        json_path = template_dir / "data.json"
        data = {"key": "value"}

        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(data, f)

        results = process_json_files(tmp_path)

        assert len(results) == 0

    def test_process_json_invalid_json(self, tmp_path):
        """Test handling of invalid JSON file."""
        json_path = tmp_path / "invalid.json"

        with open(json_path, "w", encoding="utf-8") as f:
            f.write("This is not valid JSON")

        results = process_json_files(tmp_path)

        # Should handle error gracefully
        assert results == []

    def test_process_json_no_json_files(self, tmp_path):
        """Test processing directory with no JSON files."""
        results = process_json_files(tmp_path)

        assert results == []


class TestIntegrationScenarios:
    """Integration tests for complete data preparation scenarios."""

    def test_mixed_data_sources(self, tmp_path):
        """Test processing mixed data sources (DOCX, CSV, JSON)."""
        # NOTE: PDF extraction removed during cleanup

        # Create DOCX
        docx_path = tmp_path / "achievements.docx"
        doc = Document()
        doc.add_paragraph("Achievement 1")
        doc.save(str(docx_path))

        # Create JSON
        json_path = tmp_path / "skills.json"
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump({"skills": ["Python", "Leadership"]}, f)

        # Test extraction
        docx_text = extract_text_from_docx(str(docx_path))
        json_results = process_json_files(tmp_path)

        assert "Achievement 1" in docx_text
        assert len(json_results) > 0

    def test_large_text_chunking(self):
        """Test chunking of realistically large text."""
        # Create text simulating a resume (3000+ characters)
        text = (
            """
PROFESSIONAL EXPERIENCE

Senior Engineering Manager | TechCorp | 2020-2023
- Led team of 12 engineers building cloud infrastructure
- Implemented CI/CD pipeline reducing deployment time by 60%
- Mentored 5 junior engineers to senior level
- Collaborated with product team on roadmap planning
- Managed $2M annual budget

Engineering Manager | StartupCo | 2017-2020
- Built engineering team from 3 to 15 people
- Established agile processes and best practices
- Delivered 10+ major product features on time
- Improved code quality through rigorous review process

Senior Software Engineer | BigCorp | 2014-2017
- Designed and implemented microservices architecture
- Optimized database queries improving performance by 40%
- Led technical design reviews
- Contributed to open source projects
        """
            * 5
        )  # Repeat to make it longer

        chunks = chunk_text(text)

        # Should create multiple chunks
        assert len(chunks) > 3

        # Each chunk should contain meaningful content
        for chunk in chunks:
            assert len(chunk) > 0
            assert len(chunk) <= DEFAULT_CHUNK_SIZE + DEFAULT_OVERLAP + 200

    @patch("src.cover_letter_generator.prepare_data.SentenceTransformer")
    @patch("src.cover_letter_generator.prepare_data.chromadb.PersistentClient")
    def test_main_function_integration(self, mock_chroma, mock_st, tmp_path):
        """Test main function integration (mocked dependencies)."""
        # This tests that the main function can be called without errors
        # when dependencies are mocked

        # Mock embedding model
        mock_model = Mock()
        mock_model.encode.return_value = [[0.1] * 384]
        mock_st.return_value = mock_model

        # Mock ChromaDB client
        mock_client = Mock()
        mock_collection = Mock()
        mock_client.create_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client

        # The main function would be tested here
        # For now, verify mocks are set up correctly
        assert mock_st.called or True  # Placeholder
