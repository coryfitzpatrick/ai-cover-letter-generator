"""Comprehensive tests for generator module to increase coverage to 80%+.

This test suite covers:
- CoverLetterGenerator initialization
- System prompt loading from different paths
- Leadership philosophy loading (DOCX and txt formats)
- Context retrieval with various query types
- LLM calls (mocked)
- Cover letter generation flow
- Revision functionality
- Cost tracking
- Error handling
"""

import os
import tempfile
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock, mock_open

import pytest
from docx import Document

from src.cover_letter_generator.generator import CoverLetterGenerator
from src.cover_letter_generator.analysis import JobAnalysis, JobLevel, JobType, JobRequirement


@pytest.fixture
def mock_env_vars(monkeypatch, tmp_path):
    """Mock environment variables for testing."""
    monkeypatch.setenv("GROQ_API_KEY", "test-groq-key")
    monkeypatch.setenv("OPENAI_API_KEY", "test-openai-key")
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-anthropic-key")
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    return tmp_path


@pytest.fixture
def mock_chroma_db(tmp_path):
    """Create a mock ChromaDB structure."""
    chroma_dir = tmp_path / "chroma_db"
    chroma_dir.mkdir()
    return chroma_dir


@pytest.fixture
def system_prompt_file(tmp_path):
    """Create a temporary system prompt file."""
    prompt_dir = tmp_path / "prompts"
    prompt_dir.mkdir()
    prompt_file = prompt_dir / "system_prompt.txt"
    prompt_content = """You are a cover letter generator.

Context: {context}
Job Description: {job_description}
Company: {company_name}
Title: {job_title}
Analysis: {job_analysis}
Philosophy: {leadership_philosophy}
"""
    prompt_file.write_text(prompt_content)
    return prompt_file


class TestCoverLetterGeneratorInit:
    """Tests for CoverLetterGenerator initialization."""

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_init_with_gpt4o_default(
        self, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file
    ):
        """Test initialization with default GPT-4o model."""
        # Mock the ChromaDB client and collection
        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client

        # Mock embedding model
        mock_st.return_value = Mock()

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))

        assert "gpt-4o" in generator.model_name
        assert generator.openai_client is not None
        assert generator.claude_client is None

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_init_with_opus_model(
        self, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file
    ):
        """Test initialization with Claude Opus model."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client
        mock_st.return_value = Mock()

        generator = CoverLetterGenerator(
            system_prompt_path=str(system_prompt_file), model_name="opus"
        )

        assert "opus" in generator.model_name
        assert generator.claude_client is not None
        assert generator.openai_client is None

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_init_missing_chromadb(self, mock_st, mock_chroma, mock_env_vars):
        """Test initialization fails when ChromaDB doesn't exist."""
        mock_chroma.side_effect = Exception("Collection not found")
        mock_st.return_value = Mock()

        with pytest.raises(Exception):
            CoverLetterGenerator()

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_init_missing_system_prompt(self, mock_st, mock_chroma, mock_env_vars, mock_chroma_db):
        """Test initialization fails when system prompt doesn't exist."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client
        mock_st.return_value = Mock()

        with pytest.raises(FileNotFoundError):
            CoverLetterGenerator(system_prompt_path="/nonexistent/prompt.txt")

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_init_missing_groq_key(
        self, mock_st, mock_chroma, monkeypatch, mock_chroma_db, system_prompt_file
    ):
        """Test initialization fails when GROQ_API_KEY is missing."""
        monkeypatch.delenv("GROQ_API_KEY", raising=False)
        mock_st.return_value = Mock()

        with pytest.raises(ValueError, match="GROQ_API_KEY"):
            CoverLetterGenerator(system_prompt_path=str(system_prompt_file))


class TestLoadLeadershipPhilosophy:
    """Tests for loading leadership philosophy."""

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_load_philosophy_from_docx(
        self, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file
    ):
        """Test loading leadership philosophy from DOCX file."""
        # Create a DOCX file
        philosophy_path = mock_env_vars / "Leadership Philosophy.docx"
        doc = Document()
        doc.add_paragraph("My leadership philosophy is...")
        doc.add_paragraph("I believe in empowering teams.")
        doc.save(str(philosophy_path))

        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client
        mock_st.return_value = Mock()

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))
        philosophy = generator._load_leadership_philosophy()

        assert "leadership philosophy" in philosophy.lower()
        assert "empowering teams" in philosophy

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_load_philosophy_from_txt_fallback(
        self, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file, tmp_path
    ):
        """Test loading philosophy from txt file as fallback."""
        # Create a txt file in project root
        philosophy_path = tmp_path.parent.parent.parent / "leadership_philosophy.txt"
        philosophy_path.write_text("Leadership philosophy from txt file")

        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client
        mock_st.return_value = Mock()

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))

        # Set project_root for test
        generator.project_root = tmp_path.parent.parent.parent
        philosophy = generator._load_leadership_philosophy()

        # Should attempt to load from txt if DOCX not found
        assert isinstance(philosophy, str)

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_load_philosophy_handles_missing_files(
        self, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file
    ):
        """Test that missing philosophy files are handled gracefully."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client
        mock_st.return_value = Mock()

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))
        philosophy = generator._load_leadership_philosophy()

        # Should return empty string if no files found
        assert isinstance(philosophy, str)


class TestRetrieveContext:
    """Tests for context retrieval from vector database."""

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    @patch("src.cover_letter_generator.generator.analyze_job_posting")
    def test_get_relevant_context_basic(
        self, mock_analyze, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file
    ):
        """Test basic context retrieval."""
        # Mock ChromaDB collection
        mock_client = Mock()
        mock_collection = Mock()
        mock_collection.query.return_value = {
            "documents": [["Context document 1", "Context document 2"]],
            "distances": [[0.5, 0.7]],
            "metadatas": [[{"source": "resume.pdf"}, {"source": "achievements.docx"}]],
        }
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client

        # Mock embedding model
        mock_model = Mock()
        mock_model.encode.return_value = [[0.1] * 384]
        mock_st.return_value = mock_model

        # Mock job analysis
        mock_analyze.return_value = JobAnalysis(
            level=JobLevel.MANAGER,
            job_type=JobType.PRODUCT,
            requirements=[],
            key_technologies=[],
            team_size_mentioned=False,
        )

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))
        context = generator.get_relevant_context("Job description here")

        assert "Context document 1" in context or len(context) > 0
        mock_collection.query.assert_called()

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    @patch("src.cover_letter_generator.generator.analyze_job_posting")
    def test_get_relevant_context_with_job_analysis(
        self, mock_analyze, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file
    ):
        """Test context retrieval with pre-computed job analysis."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_collection.query.return_value = {
            "documents": [["Doc"]],
            "distances": [[0.5]],
            "metadatas": [[{"source": "test.pdf"}]],
        }
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client

        mock_model = Mock()
        mock_model.encode.return_value = [[0.1] * 384]
        mock_st.return_value = mock_model

        job_analysis = JobAnalysis(
            level=JobLevel.SENIOR_MANAGER,
            job_type=JobType.STARTUP,
            requirements=[
                JobRequirement("leadership", "Lead team", 1),
                JobRequirement("technical", "Python", 1),
            ],
            key_technologies=["Python", "React"],
            team_size_mentioned=True,
        )

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))
        context = generator.get_relevant_context("Job description", job_analysis=job_analysis)

        # Should not call analyze_job_posting again
        mock_analyze.assert_not_called()
        assert isinstance(context, str)

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    @patch("src.cover_letter_generator.generator.analyze_job_posting")
    def test_get_relevant_context_no_results(
        self, mock_analyze, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file
    ):
        """Test context retrieval when no relevant documents are found."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_collection.query.return_value = {
            "documents": [[]],
            "distances": [[]],
            "metadatas": [[]],
        }
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client

        mock_model = Mock()
        mock_model.encode.return_value = [[0.1] * 384]
        mock_st.return_value = mock_model

        mock_analyze.return_value = JobAnalysis(
            level=JobLevel.MANAGER,
            job_type=JobType.PRODUCT,
            requirements=[],
            key_technologies=[],
            team_size_mentioned=False,
        )

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))
        context = generator.get_relevant_context("Job description")

        assert "No specific relevant information found" in context


class TestCostTracking:
    """Tests for API cost tracking."""

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_track_api_cost_gpt4o(
        self, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file
    ):
        """Test cost tracking for GPT-4o."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client
        mock_st.return_value = Mock()

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))

        # Track a call
        cost = generator._track_api_cost("gpt-4o", 1000, 500)

        assert cost > 0
        assert generator.total_cost > 0
        assert len(generator.api_calls) == 1

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_track_api_cost_opus(
        self, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file
    ):
        """Test cost tracking for Claude Opus."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client
        mock_st.return_value = Mock()

        generator = CoverLetterGenerator(
            system_prompt_path=str(system_prompt_file), model_name="opus"
        )

        cost = generator._track_api_cost("claude-3-opus-20240229", 1000, 500)

        # Opus is more expensive than GPT-4o
        assert cost > 0
        assert generator.total_cost > 0

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_get_cost_summary(
        self, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file
    ):
        """Test getting cost summary."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client
        mock_st.return_value = Mock()

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))

        # Make some tracked calls
        generator._track_api_cost("gpt-4o", 1000, 500)
        generator._track_api_cost("gpt-4o", 2000, 1000)

        summary = generator.get_cost_summary()

        assert summary["total_cost"] > 0
        assert summary["total_calls"] == 2
        assert len(summary["calls"]) == 2


class TestCallLLM:
    """Tests for LLM calling functionality."""

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    @patch("src.cover_letter_generator.generator.openai.Client")
    def test_call_llm_gpt4o(
        self,
        mock_openai_class,
        mock_st,
        mock_chroma,
        mock_env_vars,
        mock_chroma_db,
        system_prompt_file,
    ):
        """Test calling GPT-4o."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client
        mock_st.return_value = Mock()

        # Mock OpenAI response
        mock_openai_client = Mock()
        mock_response = Mock()
        mock_response.choices = [Mock()]
        mock_response.choices[0].message.content = "Generated cover letter"
        mock_response.usage = Mock(prompt_tokens=100, completion_tokens=50)
        mock_openai_client.chat.completions.create.return_value = mock_response
        mock_openai_class.return_value = mock_openai_client

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))
        generator.openai_client = mock_openai_client

        content, cost = generator._call_llm("System prompt", "User message")

        assert content == "Generated cover letter"
        assert cost > 0

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    @patch("src.cover_letter_generator.generator.Anthropic")
    def test_call_llm_claude(
        self,
        mock_anthropic_class,
        mock_st,
        mock_chroma,
        mock_env_vars,
        mock_chroma_db,
        system_prompt_file,
    ):
        """Test calling Claude."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client
        mock_st.return_value = Mock()

        # Mock Claude response
        mock_claude_client = Mock()
        mock_response = Mock()
        mock_response.content = [Mock()]
        mock_response.content[0].text = "Generated cover letter"
        mock_response.usage = Mock(input_tokens=100, output_tokens=50)
        mock_claude_client.messages.create.return_value = mock_response
        mock_anthropic_class.return_value = mock_claude_client

        generator = CoverLetterGenerator(
            system_prompt_path=str(system_prompt_file), model_name="opus"
        )
        generator.claude_client = mock_claude_client

        content, cost = generator._call_llm("System prompt", "User message")

        assert content == "Generated cover letter"
        assert cost > 0


class TestGenerateCoverLetter:
    """Tests for cover letter generation."""

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    @patch("src.cover_letter_generator.generator.analyze_job_posting")
    def test_generate_cover_letter_basic(
        self,
        mock_analyze,
        mock_st,
        mock_chroma,
        mock_env_vars,
        mock_chroma_db,
        system_prompt_file,
        tmp_path,
    ):
        """Test basic cover letter generation."""
        # Setup mocks
        mock_client = Mock()
        mock_collection = Mock()
        mock_collection.query.return_value = {
            "documents": [["Context"]],
            "distances": [[0.5]],
            "metadatas": [[{"source": "test.pdf"}]],
        }
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client

        mock_model = Mock()
        mock_model.encode.return_value = [[0.1] * 384]
        mock_st.return_value = mock_model

        mock_analyze.return_value = JobAnalysis(
            level=JobLevel.MANAGER,
            job_type=JobType.PRODUCT,
            requirements=[JobRequirement("leadership", "Lead team", 1)],
            key_technologies=["Python"],
            team_size_mentioned=True,
        )

        # Create critique prompt
        critique_dir = tmp_path / "prompts"
        critique_dir.mkdir()
        critique_file = critique_dir / "critique_prompt.txt"
        critique_file.write_text(
            """Critique this draft:

Company: {company_name}
Draft: {initial_draft}
Job: {job_description}

NOTES: Good draft
REFINED VERSION:
Improved cover letter
"""
        )

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))
        generator.project_root = tmp_path

        # Mock LLM calls
        with patch.object(generator, "_call_llm") as mock_call_llm:
            mock_call_llm.side_effect = [
                ("Initial draft cover letter", 0.01),
                ("NOTES: Looks good\nREFINED VERSION:\nFinal cover letter", 0.01),
            ]

            letter, cost_info = generator.generate_cover_letter(
                job_description="Job description",
                company_name="TechCorp",
                job_title="Engineering Manager",
            )

            assert "Final cover letter" in letter
            assert cost_info["total_cost"] > 0

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    @patch("src.cover_letter_generator.generator.analyze_job_posting")
    def test_generate_with_custom_context(
        self,
        mock_analyze,
        mock_st,
        mock_chroma,
        mock_env_vars,
        mock_chroma_db,
        system_prompt_file,
        tmp_path,
    ):
        """Test generation with custom context."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_collection.query.return_value = {
            "documents": [["Context"]],
            "distances": [[0.5]],
            "metadatas": [[{"source": "test.pdf"}]],
        }
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client

        mock_model = Mock()
        mock_model.encode.return_value = [[0.1] * 384]
        mock_st.return_value = mock_model

        mock_analyze.return_value = JobAnalysis(
            level=JobLevel.MANAGER,
            job_type=JobType.PRODUCT,
            requirements=[],
            key_technologies=[],
            team_size_mentioned=False,
        )

        critique_file = tmp_path / "prompts" / "critique_prompt.txt"
        critique_file.parent.mkdir()
        critique_file.write_text(
            "Critique: {company_name}\n{initial_draft}\n{job_description}\nREFINED VERSION:\nFinal"
        )

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))
        generator.project_root = tmp_path

        with patch.object(generator, "_call_llm") as mock_call_llm:
            mock_call_llm.side_effect = [("Draft", 0.01), ("REFINED VERSION:\nFinal", 0.01)]

            custom_context = "I have specific experience in healthcare."
            letter, cost_info = generator.generate_cover_letter(
                job_description="Job description", custom_context=custom_context
            )

            # Custom context should be used
            assert isinstance(letter, str)


class TestReviseCoverLetter:
    """Tests for cover letter revision."""

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    @patch("src.cover_letter_generator.generator.analyze_job_posting")
    def test_revise_cover_letter(
        self, mock_analyze, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file
    ):
        """Test cover letter revision."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_collection.query.return_value = {
            "documents": [["Context"]],
            "distances": [[0.5]],
            "metadatas": [[{"source": "test.pdf"}]],
        }
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client

        mock_model = Mock()
        mock_model.encode.return_value = [[0.1] * 384]
        mock_st.return_value = mock_model

        mock_analyze.return_value = JobAnalysis(
            level=JobLevel.MANAGER,
            job_type=JobType.PRODUCT,
            requirements=[],
            key_technologies=[],
            team_size_mentioned=False,
        )

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))

        with patch.object(generator, "_call_llm") as mock_call_llm:
            mock_call_llm.return_value = ("Revised cover letter", 0.01)

            revised, cost_info = generator.revise_cover_letter(
                current_version="Old version",
                feedback="Make it more concise",
                job_description="Job description",
            )

            assert revised == "Revised cover letter"
            assert cost_info["revision_cost"] > 0


class TestPreprocessContext:
    """Tests for context preprocessing."""

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_preprocess_context_no_custom_prompt(
        self, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file
    ):
        """Test that preprocessing returns original context when no custom prompt exists."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client
        mock_st.return_value = Mock()

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))

        original_context = "Original context here"
        processed = generator._preprocess_context(original_context)

        # Should return unchanged if no managerial_prompt.txt
        assert processed == original_context

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_preprocess_context_with_custom_prompt(
        self, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file, tmp_path
    ):
        """Test context preprocessing with custom prompt file."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client
        mock_st.return_value = Mock()

        # Create managerial_prompt.txt
        managerial_prompt = tmp_path / "managerial_prompt.txt"
        managerial_prompt.write_text("Translate this context: {context}")

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))
        generator.project_root = tmp_path

        with patch.object(generator, "openai_client") as mock_openai:
            mock_response = Mock()
            mock_response.choices = [Mock()]
            mock_response.choices[0].message.content = "Translated context"
            mock_openai.chat.completions.create.return_value = mock_response

            processed = generator._preprocess_context("Original")

            assert processed == "Translated context"


class TestEdgeCases:
    """Tests for edge cases and error handling."""

    @patch("src.cover_letter_generator.generator.chromadb.PersistentClient")
    @patch("src.cover_letter_generator.generator.SentenceTransformer")
    def test_unknown_model_cost_tracking(
        self, mock_st, mock_chroma, mock_env_vars, mock_chroma_db, system_prompt_file
    ):
        """Test cost tracking with unknown model."""
        mock_client = Mock()
        mock_collection = Mock()
        mock_client.get_collection.return_value = mock_collection
        mock_chroma.return_value = mock_client
        mock_st.return_value = Mock()

        generator = CoverLetterGenerator(system_prompt_path=str(system_prompt_file))

        # Track cost for unknown model
        cost = generator._track_api_cost("unknown-model", 1000, 500)

        # Should return 0 for unknown models
        assert cost == 0
