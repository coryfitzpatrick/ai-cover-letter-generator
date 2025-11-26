"""DOCX generation for cover letters."""

import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from dotenv import load_dotenv


def generate_cover_letter_docx(
    cover_letter_text: str,
    output_dir: Path = None,
    filename: str = None,
    contact_info: dict = None,
) -> Path:
    """Generate a cover letter DOCX file.

    Args:
        cover_letter_text: The cover letter content
        output_dir: Directory to save the DOCX (default: current directory)
        filename: Custom filename (default: cover_letter_TIMESTAMP.docx)
        contact_info: Optional dict with keys: name, email, phone, location, linkedin, portfolio

    Returns:
        Path to the generated DOCX
    """
    if output_dir is None:
        output_dir = Path.cwd()

    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"cover_letter_{timestamp}.docx"

    output_path = output_dir / filename

    # Look for template DOCX in multiple locations
    env_path = Path(__file__).parent.parent.parent / ".env"
    load_dotenv(dotenv_path=env_path)

    template_locations = []

    # 1. Check DATA_DIR/template folder (Google Drive)
    data_dir = os.getenv("DATA_DIR")
    if data_dir:
        data_dir_clean = data_dir.strip('"').strip("'")
        google_drive_template = Path(data_dir_clean).expanduser() / "template" / "Cover Letter_ AI Template.docx"
        template_locations.append(google_drive_template)

    # 2. Check project root (fallback)
    project_template = Path(__file__).parent.parent.parent / "Cover Letter_ AI Template.docx"
    template_locations.append(project_template)

    # Find first existing template
    template_path = None
    for loc in template_locations:
        if loc.exists():
            template_path = loc
            break

    # Create document from template or blank
    if template_path:
        doc = Document(str(template_path))
        # Template loaded - keep everything (header content in body or header section)
        # Find and remove only the {data} placeholder if it exists
        paragraphs_to_remove = []
        for paragraph in doc.paragraphs:
            if '{data}' in paragraph.text:
                paragraphs_to_remove.append(paragraph)

        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)
    else:
        print("Warning: DOCX template not found, creating from scratch")
        print(f"  Checked: {[str(l) for l in template_locations]}")
        doc = Document()

        # Set default font and margins if creating from scratch
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Add Header if contact info is provided
        if contact_info:
            name = contact_info.get('name', os.getenv('USER_NAME'))
            email = contact_info.get('email', '')
            phone = contact_info.get('phone', '')
            location = contact_info.get('location', '')
            linkedin = contact_info.get('linkedin', '')
            portfolio = contact_info.get('portfolio', '')

            # Name (Large, Bold) - Matches PDF NameStyle
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(name)
            name_run.bold = True
            name_run.font.name = 'Arial' # Matches Helvetica-Bold in PDF
            name_run.font.size = Pt(14) # Matches PDF size
            name_para.paragraph_format.space_after = Pt(4) # Matches PDF spaceAfter=4

            # Contact Info (Location | Phone | Email) - Matches PDF CustomHeader
            contact_parts = []
            if location:
                contact_parts.append(location)
            if phone:
                contact_parts.append(phone)
            if email:
                contact_parts.append(email)
            
            if contact_parts:
                contact_para = doc.add_paragraph(" | ".join(contact_parts))
                contact_para.style = 'Normal'
                for run in contact_para.runs:
                    run.font.name = 'Arial' # Matches PDF font
                    run.font.size = Pt(10) # Matches PDF fontSize=10
                    run.font.color.rgb = None # Default black/dark grey
                contact_para.paragraph_format.space_after = Pt(6) # Matches PDF spaceAfter=6

            # Links (LinkedIn | Portfolio) - Matches PDF CustomHeader
            link_parts = []
            if linkedin:
                link_parts.append("LinkedIn")
            if portfolio:
                link_parts.append("Portfolio")
            
            if link_parts:
                links_para = doc.add_paragraph(" | ".join(link_parts))
                links_para.style = 'Normal'
                for run in links_para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10) # Matches PDF fontSize=10
                links_para.paragraph_format.space_after = Pt(12) # Matches spacer after header

    # Add current date with tight spacing
    date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_para.style = 'Normal'
    date_run = date_para.runs[0] if date_para.runs else date_para.add_run()
    date_run.font.name = 'Arial'
    date_run.font.size = Pt(11)

    # Set tight spacing after date (0 pt)
    para_format = date_para.paragraph_format
    para_format.space_after = Pt(0)
    para_format.line_spacing = 1.0

    # Process the cover letter text
    paragraphs = cover_letter_text.strip().split('\n\n')

    for para_text in paragraphs:
        if para_text.strip():
            clean_para = para_text.strip().replace('\n', ' ')

            # Handle "Sincerely," specially with tight spacing
            if clean_para.startswith('Sincerely'):
                # Add spacing before closing
                spacer = doc.add_paragraph()
                spacer_format = spacer.paragraph_format
                spacer_format.space_after = Pt(0)
                spacer_format.line_spacing = 1.0

                # Split "Sincerely," and name for tight spacing
                if ',' in clean_para:
                    lines = clean_para.split('\n')
                    if len(lines) > 1:
                        # Already split by newline
                        sincerely_para = doc.add_paragraph(lines[0])
                        name_para = doc.add_paragraph(lines[1])
                    else:
                        # Handle "Sincerely, {Your Name}" on one line
                        parts = clean_para.split(',', 1)
                        sincerely_para = doc.add_paragraph(parts[0] + ',')
                        if len(parts) > 1:
                            name_para = doc.add_paragraph(parts[1].strip())

                    # Format "Sincerely," with tight spacing
                    sincerely_para.style = 'Normal'
                    for run in sincerely_para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                    sincerely_format = sincerely_para.paragraph_format
                    sincerely_format.space_after = Pt(0)  # Tight spacing to name
                    sincerely_format.line_spacing = 1.0

                    # Format name
                    if 'name_para' in locals():
                        name_para.style = 'Normal'
                        for run in name_para.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                        name_format = name_para.paragraph_format
                        name_format.space_after = Pt(10)
                        name_format.line_spacing = 1.0
                else:
                    # Just one paragraph
                    para = doc.add_paragraph(clean_para)
                    para.style = 'Normal'
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                    para_format = para.paragraph_format
                    para_format.space_after = Pt(10)
                    para_format.line_spacing = 1.0
            else:
                # Regular paragraph
                para = doc.add_paragraph(clean_para)
                para.style = 'Normal'

                # Set font
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)

                # Set spacing
                para_format = para.paragraph_format
                if clean_para.startswith('Dear '):
                    # Tight spacing after salutation
                    para_format.space_after = Pt(10)
                else:
                    para_format.space_after = Pt(10)
                para_format.line_spacing = 1.0

    # Save the document
    doc.save(str(output_path))

    return output_path
