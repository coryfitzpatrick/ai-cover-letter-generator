"""Prepare training data and load it into ChromaDB vector database."""

import csv
import json
import os
import re
import shutil
from pathlib import Path
from typing import List, Tuple

# Disable warnings and telemetry BEFORE importing libraries
os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["ANONYMIZED_TELEMETRY"] = "False"
os.environ["CHROMA_TELEMETRY_DISABLED"] = "True"

import chromadb
from chromadb.config import Settings
from docx import Document
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer

from .utils import get_data_directory, suppress_telemetry_errors

# Load environment variables from project root
_env_path = Path(__file__).parent.parent.parent / ".env"
load_dotenv(dotenv_path=_env_path)

# Suppress ChromaDB telemetry errors
suppress_telemetry_errors()

# Chunking configuration
DEFAULT_CHUNK_SIZE = 600  # Reduced from 1000 to better isolate achievements
DEFAULT_OVERLAP = 100     # Reduced overlap


# Embedding model
EMBEDDING_MODEL = "all-MiniLM-L6-v2"


def extract_text_from_docx(docx_path: str) -> str:
    """Extract text content from a DOCX file (Word document or exported Google Doc).

    Preserves heading information by prefixing H2 headings with [H2] marker.
    For resumes in table format, extracts table content with proper line breaks.

    Args:
        docx_path: Path to the DOCX file

    Returns:
        Extracted text content with heading markers
    """
    try:
        doc = Document(docx_path)
        text = ""

        # Extract from paragraphs first
        for paragraph in doc.paragraphs:
            # Mark H2 headings for easier company extraction
            # Check for Heading 2 style (handles variations like "Heading 2", "heading 2", etc.)
            style_name = paragraph.style.name.lower() if paragraph.style.name else ""
            if 'heading 2' in style_name or style_name == 'heading2':
                text += f"[H2]{paragraph.text}\n"
            else:
                text += paragraph.text + "\n"

        # Extract text from tables (common in resumes)
        # Process each cell's paragraphs to preserve formatting
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Extract paragraphs from cell to preserve line breaks
                    for para in cell.paragraphs:
                        cell_text = para.text.strip()
                        if cell_text:
                            text += cell_text + "\n"
                # Add extra newline between table rows
                text += "\n"

        return text.strip()
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")
        return ""


def extract_company_from_section_headers(text: str) -> str:
    """Extract company name from section headers like 'Achievements at J&J:'.

    Args:
        text: Document text to search (may include [H2] markers from DOCX headings)

    Returns:
        Company name (lowercase) or "unknown" if not found
    """
    # Look for patterns like:
    # "[H2]Achievements at J&J:"  (H2 heading)
    # "Achievements at J&J:"      (plain text)
    # "Work at Google:"
    # "Experience at Amazon:"
    # "Projects at Microsoft:"

    patterns = [
        # Pattern 1: "Company Name: XYZ" format (e.g., "[H2]Company Name: J&J" or "[H2]Company Name: J&J:")
        r'^\[H2\]\s*Company\s+Name:\s+([A-Z][A-Za-z&\s+]+?):\s*$',  # With trailing colon
        r'^\[H2\]\s*Company\s+Name:\s+([A-Z][A-Za-z&\s+]+?)\s*$',   # Without trailing colon
        # Pattern 2: H2 heading with "at Company:" format
        # Matches: "Achievements at J&J:", "Achievements at Fitbit + Google:"
        r'^\[H2\]\s*(?:Achievements|Work|Experience|Projects|Position|Impact)\s+at\s+([A-Z][A-Za-z&\s+]+?):\s*$',
        # Pattern 3: Regular line with "at Company:" format
        r'^\s*(?:Achievements|Work|Experience|Projects|Position|Impact)\s+at\s+([A-Z][A-Za-z&\s+]+?):\s*$',
        # Pattern 4: H2 heading with "Company - Type" format
        r'^\[H2\]\s*([A-Z][A-Za-z&+]+(?:\s+[A-Z][A-Za-z&+]+)*)\s*[-–—]\s*(?:Achievements|Work|Experience|Projects)',
        # Pattern 5: "Company - Type" at start of line
        r'^\s*([A-Z][A-Za-z&+]+(?:\s+[A-Z][A-Za-z&+]+)*)\s*[-–—]\s*(?:Achievements|Work|Experience|Projects)',
    ]

    for pattern in patterns:
        matches = list(re.finditer(pattern, text, re.MULTILINE))
        if matches:
            for match in matches:
                company = match.group(1).strip()
                # Validate: should be short (2-30 chars), not a full sentence
                if 2 < len(company) < 30 and not company.endswith(('.', '!', '?')):
                    # Remove trailing punctuation
                    company = company.rstrip(':').strip()
                    # Make sure it's not a common non-company word
                    if company.lower() not in ['the company', 'my company', 'this company']:
                        return company.lower()

    return "unknown"


def parse_sections_by_company(text: str) -> List[Tuple[str, str]]:
    """Parse document into sections by company based on H2 headers.

    Args:
        text: Document text with [H2] markers

    Returns:
        List of (section_text, company_name) tuples
    """
    sections = []
    current_company = "unknown"
    current_section = []

    lines = text.split('\n')

    for line in lines:
        # Check if this is an H2 header with company info
        if line.startswith('[H2]'):
            # Try to extract company from this header
            # Patterns like "[H2]Company Name: J&J" or "[H2]Achievements at Fitbit + Google:"
            patterns = [
                r'^\[H2\]\s*Company\s+Name:\s+([A-Z][A-Za-z&\s+]+?):\s*$',  # With trailing colon
                r'^\[H2\]\s*Company\s+Name:\s+([A-Z][A-Za-z&\s+]+?)\s*$',   # Without trailing colon
                r'^\[H2\]\s*(?:Achievements|Work|Experience|Projects|Position|Impact)\s+at\s+([A-Z][A-Za-z&\s+]+?):\s*$',
                r'^\[H2\]\s*([A-Z][A-Za-z&+]+(?:\s+[A-Z][A-Za-z&+]+)*)\s*[-–—]\s*(?:Achievements|Work|Experience|Projects)',
            ]

            found_company = None
            for pattern in patterns:
                match = re.match(pattern, line)
                if match:
                    company = match.group(1).strip().rstrip(':').strip()
                    # Validate company name
                    if 2 < len(company) < 30 and company.lower() not in ['the company', 'my company', 'this company']:
                        found_company = company.lower()
                        break

            # If we found a new company, save the previous section
            if found_company:
                # Save previous section if it has content
                if current_section:
                    section_text = '\n'.join(current_section).strip()
                    if section_text:
                        sections.append((section_text, current_company))
                    current_section = []
                # Update to new company
                current_company = found_company

        # Add line to current section (including the H2 header itself)
        current_section.append(line)

    # Add the final section
    if current_section:
        section_text = '\n'.join(current_section).strip()
        if section_text:
            sections.append((section_text, current_company))

    return sections


def parse_resume_sections_by_company(text: str) -> List[Tuple[str, str]]:
    """Parse resume text into sections by company.

    Resume format: Company names appear as "Company Name: Location — Title"
    followed by dates and bullet points. All content until the next company
    belongs to that company.

    Example:
        J&J MedTech: Raynham, MA — Software Tech Lead
        2023 - 2025
        • Led development team
        • Improved performance by 40%

    Args:
        text: Resume text extracted from PDF

    Returns:
        List of (section_text, company_name) tuples
    """
    sections = []
    current_company = "unknown"
    current_section = []

    lines = text.split('\n')

    # Patterns to match company headers:
    # 1. "J&J MedTech: Location — Title"
    # 2. "Google - Fitbit: Location — Title"
    company_patterns = [
        r'^([A-Z][A-Za-z&\s+-]+?):\s+[A-Z].*?—',  # Standard format with colon before location
        r'^([A-Z][A-Za-z&\s+-]+?)\s*[-–—]\s*([A-Z][A-Za-z&\s+-]+?):\s+[A-Z].*?—',  # "Company - Division:" format
    ]
    excluded_headers = ['experience', 'education', 'skills', 'summary', 'objective', 'references']

    for line in lines:
        # Check if this line starts a new company section
        match = None
        for pattern in company_patterns:
            match = re.match(pattern, line)
            if match:
                break

        if match:
            # Extract company name
            # For "Google - Fitbit:" format, combine both parts
            if match.lastindex == 2:  # Pattern matched with 2 groups
                company = f"{match.group(1).strip()} - {match.group(2).strip()}"
            else:
                company = match.group(1).strip()

            # Skip if it's a generic section header
            if company.lower() in excluded_headers:
                current_section.append(line)
                continue

            # Save previous section if it exists
            if current_section:
                section_text = '\n'.join(current_section).strip()
                if section_text:
                    sections.append((section_text, current_company))
                current_section = []

            # Clean up company name
            if 2 < len(company) < 50 and company.lower() not in ['the company', 'my company']:
                current_company = company.lower()
            else:
                current_company = "unknown"

        # Add line to current section
        current_section.append(line)

    # Add final section
    if current_section:
        section_text = '\n'.join(current_section).strip()
        if section_text:
            sections.append((section_text, current_company))

    return sections


def chunk_text(
    text: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_OVERLAP,
    metadata_header: str = ""
) -> List[str]:
    """Split text into chunks respecting paragraph boundaries.

    Args:
        text: Text to split
        chunk_size: Target size of each chunk in characters
        overlap: Number of overlapping characters between chunks
        metadata_header: Optional string to prepend to each chunk (e.g. "Source: Google (2023)\n")

    Returns:
        List of text chunks
    """
    if not text:
        return []

    # Split into paragraphs (respecting double newlines as distinct separators)
    lines = text.splitlines(keepends=True)
    
    chunks = []
    current_chunk = []
    current_length = 0
    
    # Threshold to force a split on a paragraph break
    # If we have > 200 chars and hit a blank line, we split.
    # This ensures distinct achievements (usually ~200-400 chars) are kept separate.
    SOFT_SPLIT_THRESHOLD = 200
    
    for line in lines:
        line_len = len(line)
        is_blank_line = line.strip() == ""
        
        # Check if we should force a split due to paragraph break + sufficient length
        force_split = is_blank_line and current_length > SOFT_SPLIT_THRESHOLD
        
        # If adding this line exceeds chunk size OR we are forcing a split
        if (current_length + line_len > chunk_size or force_split) and current_chunk:
            # Store current chunk
            chunk_text = "".join(current_chunk).strip()
            if chunk_text:
                if metadata_header:
                    chunk_text = f"{metadata_header}\n{chunk_text}"
                chunks.append(chunk_text)
            
            # Start new chunk
            current_chunk = []
            current_length = 0
            
            # If this was just a blank line that forced the split, we don't need to add it to the new chunk
            if is_blank_line:
                continue

        # If a single line is massive (larger than chunk size), we must split it
        if line_len > chunk_size:
            # If we have pending content, save it first
            if current_chunk:
                chunk_text = "".join(current_chunk).strip()
                if metadata_header:
                    chunk_text = f"{metadata_header}\n{chunk_text}"
                chunks.append(chunk_text)
                current_chunk = []
                current_length = 0
            
            # Split the long line using the old character-based method
            start = 0
            while start < line_len:
                end = start + chunk_size
                sub_chunk = line[start:end]
                sub_chunk = line[start:end].strip()
                if metadata_header:
                    sub_chunk = f"{metadata_header}\n{sub_chunk}"
                chunks.append(sub_chunk)
                start = end - overlap
            continue

        current_chunk.append(line)
        current_length += line_len

    # Add remaining content
    if current_chunk:
        chunk_text = "".join(current_chunk).strip()
        if chunk_text:
            if metadata_header:
                chunk_text = f"{metadata_header}\n{chunk_text}"
            chunks.append(chunk_text)

    return chunks


def process_linkedin_profile_csv(csv_path: str) -> List[Tuple[str, dict]]:
    """Extract profile information from LinkedIn Profile.csv.

    Args:
        csv_path: Path to the Profile.csv file

    Returns:
        List of tuples (text, metadata)
    """
    results = []
    try:
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                # Extract the summary/headline
                if row.get('Summary'):
                    results.append((
                        f"PROFESSIONAL SUMMARY:\n{row['Summary']}",
                        {
                            "source": "LinkedIn Profile",
                            "type": "profile_summary",
                            "name": f"{row.get('First Name', '')} {row.get('Last Name', '')}".strip()
                        }
                    ))

                if row.get('Headline'):
                    results.append((
                        f"PROFESSIONAL HEADLINE: {row['Headline']}",
                        {
                            "source": "LinkedIn Profile",
                            "type": "headline",
                            "name": f"{row.get('First Name', '')} {row.get('Last Name', '')}".strip()
                        }
                    ))
    except Exception as e:
        print(f"  Error reading Profile.csv: {e}")

    return results


def process_linkedin_recommendations_csv(csv_path: str) -> List[Tuple[str, dict]]:
    """Extract recommendations from LinkedIn Recommendations CSV.

    Args:
        csv_path: Path to the recommendations CSV file

    Returns:
        List of tuples (text, metadata)
    """
    results = []
    try:
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                if row.get('Text') and row.get('Status') == 'VISIBLE':
                    recommender = f"{row.get('First Name', '')} {row.get('Last Name', '')}".strip()
                    company = row.get('Company', 'Unknown')
                    job_title = row.get('Job Title', 'Unknown')

                    # Format the recommendation with context
                    text = (
                        f"RECOMMENDATION from {recommender} "
                        f"({job_title} at {company}):\n\n"
                        f"{row['Text']}"
                    )

                    results.append((
                        text,
                        {
                            "source": "LinkedIn Recommendation",
                            "type": "recommendation",
                            "recommender": recommender,
                            "company": company,
                            "title": job_title
                        }
                    ))
    except Exception as e:
        print(f"  Error reading recommendations CSV: {e}")

    return results


def process_csv_files(data_dir: Path) -> List[Tuple[str, dict]]:
    """Process all LinkedIn CSV files in the data directory.

    Args:
        data_dir: Path to the data directory

    Returns:
        List of tuples (text, metadata)
    """
    results = []

    # Look for CSV files in the data directory and subdirectories (exclude template folder)
    csv_files = [
        f for f in data_dir.glob("**/*.csv")
        if "template" not in str(f).lower()
    ]

    if not csv_files:
        return results

    print("\nScanning for LinkedIn CSV files...")

    for csv_file in csv_files:
        file_name = csv_file.name.lower()

        if 'profile' in file_name:
            print(f"\nProcessing {csv_file.name}...")
            profile_data = process_linkedin_profile_csv(str(csv_file))
            results.extend(profile_data)
            print(f"  Extracted {len(profile_data)} profile entries")

        elif 'recommendation' in file_name and 'received' in file_name:
            print(f"\nProcessing {csv_file.name}...")
            recommendations = process_linkedin_recommendations_csv(str(csv_file))
            results.extend(recommendations)
            print(f"  Extracted {len(recommendations)} recommendations")

    return results


def process_json_files(data_dir: Path) -> List[Tuple[str, dict]]:
    """Process JSON files in the data directory.

    Args:
        data_dir: Path to the data directory

    Returns:
        List of tuples (text, metadata)
    """
    results = []

    # Look for JSON files in the data directory and subdirectories (exclude template folder)
    json_files = [
        f for f in data_dir.glob("**/*.json")
        if "template" not in str(f).lower()
    ]

    if not json_files:
        return results

    print("\nScanning for JSON files...")

    for json_file in json_files:
        # Skip contact_info.json (it's configuration, not data)
        if json_file.name == "contact_info.json":
            continue

        print(f"\nProcessing {json_file.name}...")
        try:
            with open(json_file, "r", encoding="utf-8") as f:
                data = json.load(f)

            # Handle different JSON structures
            if isinstance(data, dict):
                # Convert dict to text entries
                for key, value in data.items():
                    if isinstance(value, (str, int, float, bool)):
                        text = f"{key}: {value}"
                        results.append((
                            text,
                            {
                                "source": json_file.name,
                                "type": "json",
                                "key": key
                            }
                        ))
                    elif isinstance(value, dict):
                        # Nested dict - convert to formatted text
                        text = f"{key}:\n" + "\n".join(
                            f"  {k}: {v}" for k, v in value.items()
                            if isinstance(v, (str, int, float, bool))
                        )
                        if text.strip():
                            results.append((
                                text,
                                {
                                    "source": json_file.name,
                                    "type": "json",
                                    "key": key
                                }
                            ))
                    elif isinstance(value, list):
                        # List of items
                        for i, item in enumerate(value):
                            if isinstance(item, str):
                                results.append((
                                    f"{key} [{i+1}]: {item}",
                                    {
                                        "source": json_file.name,
                                        "type": "json",
                                        "key": key,
                                        "index": i
                                    }
                                ))
                            elif isinstance(item, dict):
                                # List of objects
                                text = f"{key} [{i+1}]:\n" + "\n".join(
                                    f"  {k}: {v}" for k, v in item.items()
                                    if isinstance(v, (str, int, float, bool))
                                )
                                if text.strip():
                                    results.append((
                                        text,
                                        {
                                            "source": json_file.name,
                                            "type": "json",
                                            "key": key,
                                            "index": i
                                        }
                                    ))

            elif isinstance(data, list):
                # Top-level list
                for i, item in enumerate(data):
                    if isinstance(item, str):
                        results.append((
                            item,
                            {
                                "source": json_file.name,
                                "type": "json",
                                "index": i
                            }
                        ))
                    elif isinstance(item, dict):
                        text = "\n".join(
                            f"{k}: {v}" for k, v in item.items()
                            if isinstance(v, (str, int, float, bool))
                        )
                        if text.strip():
                            results.append((
                                text,
                                {
                                    "source": json_file.name,
                                    "type": "json",
                                    "index": i
                                }
                            ))

            print(f"  Extracted {len([r for r in results if r[1]['source'] == json_file.name])} entries")

        except json.JSONDecodeError as e:
            print(f"  Error parsing JSON: {e}")
        except Exception as e:
            print(f"  Error reading {json_file.name}: {e}")

    return results


def main():
    """Main function to prepare data and load into ChromaDB."""
    print("Starting data preparation...")

    # Initialize the embedding model
    print("Loading embedding model...")
    model = SentenceTransformer(EMBEDDING_MODEL)

    # Setup ChromaDB
    # Allow custom data directory via environment variable
    data_dir = get_data_directory()
    chroma_dir = data_dir / "chroma_db"

    # Delete entire chroma_db folder if it exists to clean up old data
    if chroma_dir.exists():
        print(f"Removing old ChromaDB at {chroma_dir}...")
        shutil.rmtree(chroma_dir)
        print("✓ Old database deleted")

    print(f"Initializing ChromaDB at {chroma_dir}...")
    client = chromadb.PersistentClient(
        path=str(chroma_dir),
        settings=Settings(anonymized_telemetry=False)
    )

    # Create new collection
    collection = client.create_collection(
        name="cover_letter_context",
        metadata={"description": "Context for cover letter generation"}
    )

    # Initialize document storage
    documents = []
    metadatas = []
    ids = []
    doc_id = 0

    # Process DOCX files (Word documents and exported Google Docs)
    print(f"\nScanning {data_dir} and subdirectories for DOCX files...")
    docx_files = [
        f for f in data_dir.glob("**/*.docx")
        if "template" not in str(f).lower() and not f.name.startswith("~$")  # Exclude temp files
    ]

    if not docx_files:
        print("No DOCX files found in data directory!")
    else:
        for docx_file in docx_files:
            print(f"\nProcessing {docx_file.name}...")
            text = extract_text_from_docx(str(docx_file))

            if not text:
                print("  Skipped (empty or error)")
                continue

            # Extract year from filename (applies to all chunks from this file)
            inferred_year = "unknown"
            year_match = re.search(r'20[12]\d', docx_file.name)
            if year_match:
                inferred_year = year_match.group(0)

            # Detect if this is a resume (has company sections with "Company: Location — Title" format)
            is_resume = bool(re.search(r'^[A-Z][A-Za-z&\s+]+?:\s+[A-Z].*?—', text, re.MULTILINE))

            if is_resume:
                # Parse as resume with inline company headers
                print("  Detected as Resume - parsing by company sections")
                sections = parse_resume_sections_by_company(text)
            else:
                # Parse as structured document with H2 headers
                sections = parse_sections_by_company(text)

            if not sections:
                # Fallback: treat entire document as one section with unknown company
                sections = [(text, "unknown")]

            # Track total chunks for this file
            file_chunks = []

            # Process each section separately
            for section_text, section_company in sections:
                # Create metadata header for this section
                meta_header = f"SOURCE DOCUMENT: {docx_file.name}"
                if section_company != "unknown":
                    meta_header += f"\nCOMPANY: {section_company.upper()}"
                if inferred_year != "unknown":
                    meta_header += f"\nYEAR: {inferred_year}"

                # Chunk this section
                section_chunks = chunk_text(section_text, metadata_header=meta_header)

                # Store chunks with their company
                for chunk in section_chunks:
                    file_chunks.append((chunk, section_company))

            # Count chunks by company for display
            company_counts = {}
            for _, company in file_chunks:
                company_counts[company] = company_counts.get(company, 0) + 1

            # Format output message with detailed breakdown
            print(f"  Created {len(file_chunks)} chunks total (Year: {inferred_year})")
            print(f"  Breakdown by company:")

            # Group chunks by company for display
            chunks_by_company = {}
            for chunk, company in file_chunks:
                if company not in chunks_by_company:
                    chunks_by_company[company] = []
                chunks_by_company[company].append(chunk)

            # Display each company with chunk previews
            for company in sorted(chunks_by_company.keys()):
                chunks = chunks_by_company[company]
                print(f"\n    {company.upper()} ({len(chunks)} chunks):")
                for i, chunk in enumerate(chunks, 1):
                    # Show first 150 chars of each chunk (after metadata header)
                    lines = chunk.split('\n')
                    # Skip metadata lines (SOURCE DOCUMENT, COMPANY, YEAR)
                    content_lines = [l for l in lines if not l.startswith('SOURCE DOCUMENT:')
                                                      and not l.startswith('COMPANY:')
                                                      and not l.startswith('YEAR:')]
                    content = '\n'.join(content_lines).strip()
                    preview = content[:150].replace('\n', ' ')
                    print(f"      [{i}] {preview}...")
            print()  # Empty line after all companies

            # Add each chunk as a document
            for i, (chunk, chunk_company) in enumerate(file_chunks):
                documents.append(chunk)
                # Use relative path from data_dir for source
                relative_path = docx_file.relative_to(data_dir)
                metadatas.append({
                    "source": str(relative_path),
                    "type": "docx",
                    "chunk_index": i,
                    "total_chunks": len(file_chunks),
                    "company": chunk_company,
                    "year": inferred_year
                })
                ids.append(f"doc_{doc_id}")
                doc_id += 1

    # Process LinkedIn CSV files
    csv_data = process_csv_files(data_dir)

    if csv_data:
        for text, metadata in csv_data:
            documents.append(text)
            metadatas.append(metadata)
            ids.append(f"doc_{len(documents) - 1}")

    # Process JSON files
    json_data = process_json_files(data_dir)

    if json_data:
        for text, metadata in json_data:
            documents.append(text)
            metadatas.append(metadata)
            ids.append(f"doc_{len(documents) - 1}")

    if not documents:
        print("\nNo documents to process!")
        return

    # Generate embeddings and add to collection
    print(f"\nGenerating embeddings for {len(documents)} document chunks...")
    embeddings = model.encode(documents, show_progress_bar=True)

    print("Adding documents to ChromaDB...")
    collection.add(
        embeddings=embeddings.tolist(),
        documents=documents,
        metadatas=metadatas,
        ids=ids
    )

    # Count file types
    docx_count = len(docx_files) if docx_files else 0
    csv_count = len([
        m for m in metadatas
        if m.get('type') in ['profile_summary', 'headline', 'recommendation']
    ])

    print(f"\n✓ Successfully processed {docx_count} DOCX files (Word/Google Docs)")
    if csv_count > 0:
        print(f"✓ Processed {csv_count} LinkedIn entries (profile + recommendations)")
    print(f"✓ Created {len(documents)} total document chunks")
    print(f"✓ Database saved at: {chroma_dir}")


if __name__ == "__main__":
    main()
